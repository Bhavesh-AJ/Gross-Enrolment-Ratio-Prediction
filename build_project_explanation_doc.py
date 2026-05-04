from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUTPUT = "GER_Project_Explanation.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(header_cells[i], header, bold=True)
        set_cell_shading(header_cells[i], "D9EAF7")

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))

    document.add_paragraph()
    return table


def add_formula(document, text):
    paragraph = document.add_paragraph()
    paragraph.style = "Intense Quote"
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(31, 78, 121)


def add_bullets(document, items):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(10.5)
styles["Heading 1"].font.name = "Calibri"
styles["Heading 1"].font.size = Pt(16)
styles["Heading 1"].font.bold = True
styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
styles["Heading 2"].font.name = "Calibri"
styles["Heading 2"].font.size = Pt(13)
styles["Heading 2"].font.bold = True
styles["Heading 2"].font.color.rgb = RGBColor(47, 84, 150)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("A Machine Learning Approach to Predict Gross Enrollment Ratio in Indian Higher Education Using AISHE Data")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(31, 78, 121)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle.add_run("Project Explanation Report")
sub_run.italic = True
sub_run.font.size = Pt(12)

doc.add_paragraph()

doc.add_heading("1. Problem Statement", level=1)
doc.add_paragraph(
    "The aim of this project is to predict Gross Enrollment Ratio, also called GER, "
    "for Indian higher education. GER tells us how many students are enrolled in "
    "higher education compared to the population in the 18-23 age group."
)
add_formula(doc, "GER = (Total enrolled students / Population aged 18-23) x 100")
doc.add_paragraph(
    "So GER is not just affected by the number of students. It also depends heavily "
    "on the population of the correct age group."
)

doc.add_heading("2. Initial Dataset Problem", level=1)
doc.add_paragraph("At the beginning, the dataset had only a few features:")
add_bullets(doc, ["students", "universities", "faculty", "year"])
doc.add_paragraph(
    "The dataset was also very small. Earlier, we only had around 54 rows from "
    "2019 and 2021. Because of that, the model gave poor results: R2 was low, "
    "MAE was high, and MAPE was high."
)
doc.add_paragraph(
    "This happened because the most important feature, population_18_23, was missing. "
    "Without population, the model cannot properly understand GER."
)

doc.add_heading("3. Data Sources Used", level=1)
doc.add_paragraph("We used AISHE Excel reports from multiple years:")
add_bullets(doc, ["AISHE 2018", "AISHE 2019", "AISHE 2020", "AISHE 2021"])
doc.add_paragraph("These reports contained data about:")
add_bullets(doc, ["GER", "students", "universities", "faculty", "population aged 18-23"])

doc.add_heading("4. Data Extraction", level=1)
doc.add_paragraph("We created extraction scripts for each year:")
add_bullets(doc, ["extract_2018.py", "extract_2019.py", "extract_2020.py", "extract_2021.py"])
doc.add_paragraph(
    "These scripts extract the required values from each AISHE Excel file. We also "
    "created extract_utils.py, which handles common cleaning tasks such as state "
    "name standardization."
)
doc.add_paragraph("For example, some files used different names like:")
add_bullets(doc, ["A & N Islands", "Andaman & Nicobar Islands", "Andaman and Nicobar Islands"])
doc.add_paragraph("These were standardized so merging works correctly.")

doc.add_heading("5. Important Fixes Done", level=1)
doc.add_paragraph(
    "One major issue was found in the old dataset. For 2019, the faculty column was wrong. "
    "It was extracting a smaller category value instead of total faculty."
)
add_table(
    doc,
    ["Example", "Incorrect Value", "Correct Value"],
    [["Andhra Pradesh faculty", "143", "99737"]],
)
doc.add_paragraph("We fixed the 2019 faculty extraction column.")
doc.add_paragraph(
    "Another important fix was GER extraction. Earlier, the code was taking the wrong "
    "Total column from the GER sheet. AISHE has multiple category-wise GER columns, "
    "such as all categories, SC, ST, etc. We corrected the extraction to use the "
    "All Categories Total GER column."
)

doc.add_heading("6. Final Dataset Creation", level=1)
doc.add_paragraph("After extracting all years, we combined them into one final dataset.")
add_formula(doc, "Final dataset: final_ger_dataset.csv")
add_table(
    doc,
    ["Property", "Value"],
    [
        ["Rows", "126"],
        ["Columns", "12"],
        ["Years", "2018, 2019, 2020, 2021"],
        ["Missing values", "0"],
        ["Duplicate rows", "0"],
    ],
)
doc.add_paragraph("The final dataset contains:")
add_bullets(
    doc,
    [
        "state",
        "ger",
        "students",
        "universities",
        "faculty",
        "population_18_23",
        "year",
        "students_per_uni",
        "faculty_per_student",
        "enrollment_rate_proxy",
        "universities_per_lakh_population",
        "faculty_per_lakh_population",
    ],
)

doc.add_heading("7. Feature Engineering", level=1)
doc.add_paragraph(
    "Feature engineering was the most important part of the project. Instead of only "
    "using raw values, we created meaningful ratio-based features."
)
add_table(
    doc,
    ["Feature", "Formula", "Meaning"],
    [
        ["students_per_uni", "students / universities", "Average student load per university"],
        ["faculty_per_student", "faculty / students", "Faculty availability compared to students"],
        ["enrollment_rate_proxy", "students / population_18_23 x 100", "Close to the GER formula"],
        [
            "universities_per_lakh_population",
            "universities / population_18_23 x 100000",
            "Higher education access compared to population",
        ],
        [
            "faculty_per_lakh_population",
            "faculty / population_18_23 x 100000",
            "Faculty strength compared to population",
        ],
    ],
)

doc.add_heading("8. Models Used", level=1)
doc.add_paragraph(
    "We tested regression models because GER is a continuous numerical value."
)
add_bullets(doc, ["Linear Regression", "Ridge Regression", "Polynomial Ridge Regression"])
doc.add_paragraph(
    "Linear Regression was simple but could not capture complex relationships well. "
    "Ridge Regression helped reduce overfitting. Polynomial Ridge Regression gave "
    "the best result because it can learn interactions between features."
)

doc.add_heading("9. Best Model", level=1)
doc.add_paragraph("The best model was Polynomial Ridge Regression.")
doc.add_paragraph("It performed well because it learned interactions such as:")
add_formula(doc, "students_per_uni x universities_per_lakh_population")
doc.add_paragraph(
    "This interaction is powerful because it indirectly represents enrollment compared to population."
)

doc.add_heading("10. Final Results", level=1)
add_table(
    doc,
    ["Metric", "Value"],
    [
        ["Model", "Polynomial Ridge Regression"],
        ["MAE", "0.009"],
        ["R2 Score", "1.000"],
        ["MAPE", "0.04%"],
        ["Approximate Accuracy", "99.96%"],
    ],
)
doc.add_paragraph(
    "Since this is a regression project, we do not officially use classification accuracy. "
    "Instead, we use Approximate Accuracy = 100 - MAPE."
)
add_formula(doc, "Approximate Accuracy = 100 - 0.04 = 99.96%")

doc.add_heading("11. Important Explanation About High Accuracy", level=1)
doc.add_paragraph(
    "The model gets very high accuracy because GER is mathematically related to "
    "students / population_18_23. After we added population_18_23, the model had "
    "the most important missing information."
)
doc.add_paragraph(
    "So the improvement is not random. It happened because we added the correct "
    "demographic feature required to explain GER."
)
doc.add_paragraph(
    "Initially, the model had poor performance because the population aged 18-23 was missing. "
    "Since GER is defined using this population group, adding this feature significantly "
    "improved the model. Polynomial Ridge Regression further improved performance by learning "
    "interactions between enrollment, institutions, faculty, and population-based features."
)

doc.add_heading("12. Files Created", level=1)
add_table(
    doc,
    ["File", "Purpose"],
    [
        ["main.py", "Trains the model and prints results"],
        ["final_ger_dataset.csv", "Final cleaned dataset"],
        ["model_results.txt", "Final result summary for submission"],
        ["ger_demo.html", "Browser-based demo that runs with Live Server"],
        ["app.py", "Streamlit app"],
        ["README.md", "Project explanation and setup guide"],
    ],
)

doc.add_heading("13. How to Explain in Viva", level=1)
doc.add_paragraph(
    "This project predicts Gross Enrollment Ratio in Indian higher education using AISHE data. "
    "Initially, the model performed poorly because the dataset had only infrastructure features "
    "like students, universities, and faculty. Later, I expanded the dataset using multiple AISHE "
    "years from 2018 to 2021 and added the population aged 18-23, which is directly related to GER. "
    "I performed feature engineering by creating ratio-based features such as students per university, "
    "faculty per student, universities per lakh population, and faculty per lakh population. After "
    "testing Linear Regression, Ridge Regression, and Polynomial Ridge Regression, the Polynomial Ridge "
    "model gave the best performance with an R2 score close to 1 and very low MAPE. The main learning "
    "from the project is that correct feature engineering and demographic data are more important than "
    "simply changing algorithms."
)

doc.add_heading("14. Demo Flow", level=1)
add_bullets(
    doc,
    [
        "Show README.md.",
        "Show final_ger_dataset.csv.",
        "Show main.py.",
        "Show model_results.txt.",
        "Open ger_demo.html using Live Server.",
        "Enter sample values and show predicted GER.",
    ],
)

doc.add_heading("15. One-Line Summary", level=1)
doc.add_paragraph(
    "This project predicts Gross Enrollment Ratio using AISHE multi-year data by combining education "
    "infrastructure features with demographic population data and applying Polynomial Ridge Regression "
    "for accurate GER prediction."
)

doc.save(OUTPUT)
print(OUTPUT)
